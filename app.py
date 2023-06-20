import PySimpleGUI as sg
import math
import random
import docx
import hashlib
from docx import Document
import chardet
from colorama import init
init()


class AppRSA:
    # code
    # step 1 : Tạo Key
    # gcd
    def gcd(self, a, b):
        if b == 0:
            return a
        return self.gcd(b, a % b)

    def is_prime(self, n):
        if n <= 1:
            return False
        for i in range(2, math.isqrt(n) + 1):
            if n % i == 0:
                return False
        return True

    def random_prime(self):
        while True:
            a = random.randint(3, 150)
            if self.is_prime(a):
                return a

    def phi_n(self, p, q):
        if self.is_prime(p) and self.is_prime(q):
            phi = (p - 1) * (q - 1)
            return phi



    # def euclid(self, e, phi):
    #     x1, x2, y1, y2 = 1, 0, 0, 1
    #     a, b = e, phi
    #     while b != 0:
    #         q = a // b
    #         a, b = b, a - q * b
    #         x1, x2 = x2, x1 - q * x2
    #         y1, y2 = y2, y1 - q * y2
    #     if a == 1:
    #         return x1 % phi
    #     else:
    #         raise ValueError("The multiplicative inverse does not exist")
    def euclid(self, b, phi_n):
        if self.gcd(b, phi_n) != 1:
            return None
        m0 = phi_n
        y = 0
        x = 1
        if phi_n == 1:
            return 0
        while b > 1:
            q = b // phi_n
            t = phi_n
            phi_n = b % phi_n
            b = t
            t = y
            y = x - q * y
            x = t
        if x < 0:
            x += m0
        return x

    def public_key(self, p, q):
        # return 61
        phi = self.phi_n(p, q)
        arr = []
        for i in range(2, phi):
            if (self.gcd(i, phi) == 1) and (self.is_prime(i)):
                arr.append(i)
        if len(arr) == 0:
            return None
        return random.choice(arr)

    def private_key(self, p, q, e):
        phi = self.phi_n(p, q)
        d = self.euclid(e, phi)
        return d

    def key(self, values):
        p = int(values[self.P_key])
        q = int(values[self.Q_key])
        e = self.public_key(p, q)
        d = self.private_key(p, q, e)
        return e, d

    def read_file(self, file_path):

        if file_path.endswith('.txt'):
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected_result = chardet.detect(raw_data)
                detected_encoding = detected_result['encoding']
                file_content = raw_data.decode(detected_encoding)
                return file_content

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            paragraphs = doc.paragraphs

            formatted_text = []
            for paragraph in paragraphs:
                text = paragraph.text
                formatted_text.append(text)
            return '\n'.join(formatted_text)
        else:
            return "Unsupported file format."

    def rgb_to_hex(self, r, g, b):
        return "#{:02x}{:02x}{:02x}".format(r, g, b)

    def read_text_color(self, file_path):
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            paragraphs = doc.paragraphs

            for paragraph in paragraphs:
                runs = paragraph.runs

                for run in runs:
                    if run.font.color.rgb is not None:
                        r, g, b = run.font.color.rgb
                        hex_color = self.rgb_to_hex(r, g, b)

                        # Áp dụng màu chữ cho đoạn văn bản
                        colored_text = f"\033[38;2;{r};{g};{b}m{run.text}\033[0m"
                        print(colored_text)
        else:
            print("Unsupported file format.")

    def save_data_to_file(self, file_path, data):
        if file_path.endswith('.txt'):
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(data)
            print(f"File saved successfully: {file_path}")
        elif file_path.endswith('.docx'):
            doc = docx.Document()
            doc.add_paragraph(data)
            doc.save(file_path)
            print(f"File saved successfully: {file_path}")
        else:
            print("Invalid file format. Please select a valid file type.")

    def MD5(self, mess):
        mess_bytes = mess.encode('utf-8')
        # Tính toán giá trị băm
        hash_object = hashlib.md5(mess_bytes)
        toHex = hash_object.hexdigest()
        return toHex

    def PowMod(self, base, exponent, modulus):
        if exponent == 0:
            return 1

        result = 1
        baseValue = base % modulus
        exp = exponent

        while exp > 0:
            if exp % 2 == 1:
                result = (result * baseValue) % modulus

            baseValue = (baseValue * baseValue) % modulus
            exp = exp // 2

        if result < 0:
            result = (result + modulus) % modulus

        return result

    def decimal(self, hex):
        res = ""
        for val in hex:
            value = int(val, 16)
            res = res + str(value) + '-'
        return res[:-1]

    def chu_ky(self, text, p, q, values):
        res = self.MD5(text)
        dcm = self.decimal(res)
        elements = dcm.split('-')
        a = []
        s = ""
        for el in elements:
            if '-' not in el:
                val = self.PowMod(int(el), int(values[self.PrivateKey]), p * q)
                a.append(val)
        for i in range(len(a)):
            s += str(a[i]) + '-'
        return s

    def Ky(self, values):
        p = int(values[self.P_key])
        q = int(values[self.Q_key])
        mess = values['-Input-']
        messMD5 = self.MD5(mess)
        messChuKy = self.chu_ky(mess, p, q, values)
        self.window['-MD5-'].update(messMD5)
        self.window[self.OutputText_key].update(messChuKy)

    def check_chu_ky(self, values):
        tmp = self.MD5(values[self.Input])
        res = values[self.Output].split("-")
        p = int(values[self.P_key])
        q = int(values[self.Q_key])
        d = int(values[self.PublicKey])
        n = p * q
        new_arr_sign = [self.PowMod(int(item), d, n) for item in res if item]
        print(new_arr_sign)
        string = ""
        for val in new_arr_sign:
            # value = self.decimal_to_hex(val)
            value = str(hex(val))[2:]
            string += value
        print(tmp, string)
        if tmp == string:
            sg.popup("Chữ ký đúng", title="Thông báo")
        else:
            sg.popup("Chữ ký sai", title="Thông báo")

    # UI
    def __init__(self):
        sg.set_options(font=("Arial Bold", 14))
        self.P_key = '-P-'
        self.Q_key = '-Q-'
        self.PublicKey = '-PublicKey-'
        self.PrivateKey = '-PrivateKey-'
        self.Input_key = '-Input-'
        self.MD5_key = '-MD5-'
        self.OutputText_key = '-OutputText-'
        self.Output = 'OutputText'
        self.Input = 'Input'
        self.Primary_Color = "#00BFFF"

        self.layout = [
            # col1
            [
                sg.Column([
                    # row 1
                    [
                        sg.Column([
                            [sg.Text("Tạo Khóa", background_color=self.Primary_Color)]
                        ], justification='center', background_color=self.Primary_Color)
                    ],
                    [sg.VerticalSeparator("None")],

                    [
                        sg.Column([
                            # row 2
                            [sg.Text("P:", background_color=self.Primary_Color),
                             sg.Input(key=self.P_key, size=(11, None)),
                             sg.VerticalSeparator(pad=((2, 2), (1, 1)), color='red'),
                             sg.Text("Q:", background_color=self.Primary_Color),
                             sg.Input(key=self.Q_key, size=(10, None))
                             ]
                        ], background_color=self.Primary_Color)
                    ],
                    [sg.VerticalSeparator("None")],
                    # row 4
                    [sg.Button('Public key', size=(10, 1), font=("Arial Bold", 10)),
                     sg.Text("e:", font=("Arial Bold", 12), background_color=self.Primary_Color),
                     sg.Input(key=self.PublicKey, readonly=True, size=(7, None)),
                     sg.Text("n:", font=("Arial Bold", 12), background_color=self.Primary_Color),
                     sg.Input(key='-n-', readonly=True, size=(7, None)),
                     ],
                    # row 5
                    [sg.Button('Private key', size=(10, 1), button_color='red', font=("Arial Bold", 10)),
                     sg.Text("d:", font=("Arial Bold", 12), background_color=self.Primary_Color),
                     sg.Input(key=self.PrivateKey, readonly=True, size=(7, None)),
                     sg.Text("n:", font=("Arial Bold", 12), background_color=self.Primary_Color),
                     sg.Input(key='--n--', readonly=True, size=(7, None)),
                     ],
                    [sg.VerticalSeparator("None")],
                    [
                        sg.Column([
                            [sg.Button('Tạo', button_color='blue', size=(10, 1), font=("Arial Bold", 12)),

                             sg.Button('Reset', button_color="red", size=(10, 1), font=("Arial Bold", 12)),
                             ]
                        ], justification="center", background_color=self.Primary_Color)
                    ],
                    # row 6
                    # [sg.Button('Tạo', button_color='blue', font=("Arial Bold", 12), size=(10, 1)),
                    #  sg.Button('Reset', button_color="red", font=("Arial Bold", 12), size=(10, 1))
                    #  ],
                    [
                        sg.Column([
                            [sg.Button('Random(P, Q)', button_color='blue', font=("Arial Bold", 12), size=(10, 1))]
                        ], justification='center', background_color=self.Primary_Color)
                    ]
                ], vertical_alignment='top', background_color=self.Primary_Color),

                sg.VerticalSeparator(pad=((20, 20), (0, 0))),
                # col 2

                sg.Column([
                    [sg.Text("Văn bản cần ký", background_color=self.Primary_Color, key='Original')],
                    [sg.VerticalSeparator("None")],
                    [sg.Multiline('', size=(20, 5), key=self.Input_key)],
                    [
                        sg.Column([
                            [sg.Button('File', button_color='blue', size=(5, 1), font=("Arial Bold", 12)),
                             sg.Button('Ký', size=(5, 1), font=("Arial Bold", 12)),
                             ]
                        ], justification="center", background_color=self.Primary_Color)
                    ],
                    [sg.VerticalSeparator("None")],
                    [sg.Text("Hàm băm", background_color=self.Primary_Color)],
                    [sg.Multiline('', size=(20, 3), key=self.MD5_key)],
                    [sg.VerticalSeparator("None")],
                    [sg.Text("Chữ ký", background_color=self.Primary_Color)],
                    [sg.Multiline('', size=(20, 5), key=self.OutputText_key)],
                    [
                        sg.Column([
                            [sg.Button('Chuyển', button_color='blue', size=(10, 1), font=("Arial Bold", 12)),
                             sg.Button('Save', button_color='blue', size=(10, 1), font=("Arial Bold", 12))]
                        ], justification='center', background_color=self.Primary_Color)
                    ],

                ], background_color=self.Primary_Color),
                sg.VerticalSeparator(pad=((20, 20), (0, 0))),

                # col 3
                sg.Column([
                    [sg.Text("Văn bản ký", background_color=self.Primary_Color)],
                    [sg.VerticalSeparator("None")],
                    [sg.Multiline('', size=(20, 5), key=self.Input)],
                    [
                        sg.Column([
                            [sg.Button('File văn bản', button_color='blue', size=(10, 1), font=("Arial Bold", 12))]
                        ], justification='center', background_color=self.Primary_Color)
                    ],
                    [sg.VerticalSeparator("None")],
                    [sg.Text("Chữ ký", background_color=self.Primary_Color)],
                    [sg.Multiline('', size=(20, 5), key=self.Output)],
                    [
                        sg.Column([
                            [sg.Button('File chữ ký', button_color='blue', size=(10, 1), font=("Arial Bold", 12))],
                        ], justification='center', background_color=self.Primary_Color)
                    ],
                    [
                        sg.Column([
                            [sg.Button('Kiểm tra', button_color='#20B2AA', font=("Arial Bold", 12), size=(10, 1))]
                        ], justification='center', background_color=self.Primary_Color)
                    ]
                ], vertical_alignment='top', background_color=self.Primary_Color),
            ],
        ]
        self.window = sg.Window('RSA Digital Signature', self.layout, size=(1000, 550),
                                background_color=self.Primary_Color)
        self.is_running = True
        self.default_values = {
            self.P_key: "",
            self.Q_key: "",
            self.PublicKey: "",
            self.PrivateKey: "",
            self.Input_key: "",
            self.OutputText_key: "",
            self.Input: "",
            self.Output: "",
            self.MD5_key: "",
        }

    def run(self):
        while self.is_running:
            event, values = self.window.read()
            if event == sg.WINDOW_CLOSED:
                self.is_running = False
            elif event == 'Tạo':
                p = int(values[self.P_key])
                q = int(values[self.Q_key])
                # p, q = self.random_prime()
                if p <= 1 or q <= 1:
                    sg.popup('P, Q phải lơn hơn 1')
                    continue
                if not self.is_prime(p) or not self.is_prime(q):
                    sg.popup('P và Q phải là số nguyên tố')
                    continue
                # if type(p) is str or type(q) is str:
                #     sg.popup('P and Q is int')
                e = self.public_key(p, q)
                d = self.private_key(p, q, e)
                n = p * q
                public = e
                private = d
                self.window[self.PublicKey].update(public)
                self.window['-n-'].update(n)
                self.window[self.PrivateKey].update(private)
                self.window['--n--'].update(n)
                print(public)
                print(private)
            elif event == 'Reset':
                for key in self.default_values:
                    self.window[key].update(self.default_values[key])
            elif event == 'Random(P, Q)':
                p = self.random_prime()
                q = self.random_prime()
                self.window[self.P_key].update(p)
                self.window[self.Q_key].update(q)
            elif event == 'File':
                file_path = sg.popup_get_file('Select a file',
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    file_content = self.read_file(file_path)
                    # colored_content = self.read_text_color(file_path)
                    print(self.read_text_color(file_path))
                    self.window['-Input-'].update(file_content)
                    # self.window['-Input-'].update(value=colored_content, background_color='#FFFFFF',
                    #                                text_color=self.read_text_color(file_path))
            elif event == "Ký":
                self.Ky(values)
            elif event == "Chuyển":
                res = values[self.Input_key]
                tmp = values[self.OutputText_key]
                self.window['Input'].update(res)
                self.window['OutputText'].update(tmp)
            elif event == "File văn bản":
                file_path = sg.popup_get_file('Select a file',
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    file_content = self.read_file(file_path)
                    self.window['Input'].update(file_content)
            elif event == "File chữ ký":
                file_path = sg.popup_get_file('Select a file',
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    file_content = self.read_file(file_path)
                    self.window[self.Output].update(file_content)
            elif event == 'Save':

                file_path = sg.popup_get_file('Save File', save_as=True,
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    data = values[self.OutputText_key]

                    rsa.save_data_to_file(file_path, data)
            elif event == "Kiểm tra":
                self.check_chu_ky(values)
        self.window.close()


rsa = AppRSA()
rsa.run()


