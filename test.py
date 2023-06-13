import PySimpleGUI as sg
import math
import random
import docx
import hashlib
from docx import Document
import docx2txt
import chardet
from unidecode import unidecode
from docx.shared import RGBColor


class AppRSA:
    # code
    # step 1 : Tạo Key
    # gcd
    def gcd(self, a, b):
        if b == 0:
            return a
        return self.gcd(b, a % b)

    def is_prime(self, n):
        if n <= 1:
            return False
        for i in range(2, math.isqrt(n) + 1):
            if n % i == 0:
                return False
        return True

    def phi_n(self, p, q):
        if self.is_prime(p) and self.is_prime(q):
            phi = (p - 1) * (q - 1)
            return phi

    def public_key(self, p, q):
        phi = self.phi_n(p, q)
        while True:
            b = random.randrange(2, phi)
            if self.gcd(b, phi) == 1:
                return b

    def euclid(self, e, phi):
        x1, x2, y1, y2 = 1, 0, 0, 1
        a, b = e, phi
        while b != 0:
            q = a // b
            a, b = b, a - q * b
            x1, x2 = x2, x1 - q * x2
            y1, y2 = y2, y1 - q * y2
        if a == 1:
            return x1 % phi
        else:
            raise ValueError("The multiplicative inverse does not exist")

    def private_key(self, p, q):
        phi = self.phi_n(p, q)
        e = self.public_key(p, q)
        d = self.euclid(e, phi)
        return d

    def read_file(self, file_path):

        if file_path.endswith('.txt'):
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected_result = chardet.detect(raw_data)
                detected_encoding = detected_result['encoding']
                file_content = raw_data.decode(detected_encoding)
                return file_content

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            paragraphs = doc.paragraphs

            formatted_text = []
            for paragraph in paragraphs:
                text = paragraph.text
                formatted_text.append(text)

            return '\n'.join(formatted_text)
        else:
            return "Unsupported file format."

    def save_data_to_file(self, file_path, data):
        if file_path.endswith('.txt'):
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(data)
            print(f"File saved successfully: {file_path}")
        elif file_path.endswith('.docx'):
            doc = docx.Document()
            doc.add_paragraph(data)
            doc.save(file_path)
            print(f"File saved successfully: {file_path}")
        else:
            print("Invalid file format. Please select a valid file type.")

    def MD5(self, mess):
        mess_bytes = mess.encode('utf-8')
        # Tính toán giá trị băm
        hash_object = hashlib.md5(mess_bytes)
        toHex = hash_object.hexdigest()
        return toHex

    def PowMod(self, base, exponent, modulus):
        if exponent == 0:
            return 1

        result = 1
        baseValue = base % modulus
        exp = exponent

        while exp > 0:
            if exp % 2 == 1:
                result = (result * baseValue) % modulus

            baseValue = (baseValue * baseValue) % modulus
            exp = exp // 2

        if result < 0:
            result = (result + modulus) % modulus

        return result

    def decimal(self, hex):
        res = ""
        for val in hex:
            value = int(val, 16)
            res = res + str(value) + '-'
        return res[:-1]

    def chu_ky(self, text, p, q):
        res = self.MD5(text)
        dcm = self.decimal(res)
        elements = dcm.split('-')
        a = []
        s = ""
        for el in elements:
            if '-' not in el:
                val = self.PowMod(int(el), self.public_key(p, q), p * q)
                a.append(val)
        for i in range(len(a)):
            s += str(a[i]) + '-'
        return s

    def check_chu_ky(self, values):
        # inputStr = values[self.Input]
        # arrSignature = (values[self.Output]).split('-')
        # hashCheck = self.MD5(inputStr)
        # newArr = []
        # for item in arrSignature:
        #     try:
        #         num = int(item)
        #         newArr.append(self.PowMod(num, self.public_key(p, q), p * q))
        #     except ValueError:
        #         print("Giá trị không hợp lệ:", item)
        # checkHash = [self.decimal(item) for item in newArr]
        # checkHashVerify = "".join(checkHash)
        # check = hashCheck == checkHashVerify
        # if check:
        #     sg.popup("Chữ ký đúng")
        # else:
        #     sg.popup("Chữ ký đúng")
        tmp = values[self.Input]
        res = values[self.MD5_key]
        if res == self.MD5(tmp):
            sg.popup("chữ kí đúng")
        else:
            sg.popup("Chữ ký sai")

    def Ky(self, values):
        p = int(values[self.P_key])
        q = int(values[self.Q_key])
        mess = values['-Input-']
        messToHex = self.MD5(mess)
        messChuKy = self.chu_ky(mess, p, q)
        self.window['-MD5-'].update(messToHex)
        self.window[self.OutputText_key].update(messChuKy)

    # UI
    def __init__(self):
        sg.set_options(font=("Arial Bold", 14))
        self.P_key = '-P-'
        self.Q_key = '-Q-'
        self.PublicKey = '-PublicKey-'
        self.PrivateKey = '-PrivateKey-'
        self.Input_key = '-Input-'
        self.MD5_key = '-MD5-'
        self.OutputText_key = '-OutputText-'
        self.Output = 'OutputText'
        self.Input = 'Input'
        self.Primary_Color = "#00BFFF"
        self.layout = [
            # col1
            [
                sg.Column([
                    # row 1
                    [sg.Text("Tạo Khóa", background_color=self.Primary_Color)],
                    [sg.VerticalSeparator("None")],
                    # row 2
                    [sg.Text("P:", size=(7, 1), justification='center'), sg.Input(key=self.P_key, size=(15, 1))],
                    # row 3
                    [sg.Text("Q:", size=(7, 1), justification='center'), sg.Input(key=self.Q_key, size=(15, 1))],
                    [sg.VerticalSeparator("None")],
                    # row 4
                    [sg.Button('Public key', size=(10, 1), font=("Arial Bold", 10)),
                     sg.Input(key=self.PublicKey, readonly=True, size=(15, None))
                     ],
                    # row 5
                    [sg.Button('Private key', size=(10, 1), button_color='red', font=("Arial Bold", 10)),
                     sg.Input(key=self.PrivateKey, readonly=True, size=(15, None))],
                    [sg.VerticalSeparator("None")],
                    # row 6
                    [sg.Button('Tạo', button_color='blue', font=("Arial Bold", 12), size=(10, 1)),
                     sg.Button('Reset', button_color="red", font=("Arial Bold", 12), size=(10, 1))
                     ],
                ], vertical_alignment='top', background_color=self.Primary_Color),

                sg.VerticalSeparator(pad=((20, 20), (0, 0))),
                # col 2

                sg.Column([
                    [sg.Text("Văn bản cần ký", background_color=self.Primary_Color, key='Original')],
                    [sg.VerticalSeparator("None")],
                    [sg.Multiline('', size=(20, 5), key=self.Input_key)],
                    [
                        sg.Column([
                            [sg.Button('File', button_color='blue', size=(5, 1), font=("Arial Bold", 12)),
                             sg.Button('Ký', size=(5, 1), font=("Arial Bold", 12)),
                             ]
                        ], justification="center", background_color=self.Primary_Color)
                    ],
                    [sg.VerticalSeparator("None")],
                    [sg.Text("Hàm băm", background_color=self.Primary_Color)],
                    [sg.Multiline('', size=(20, 3), key=self.MD5_key)],
                    [sg.VerticalSeparator("None")],
                    [sg.Text("Chữ ký", background_color=self.Primary_Color)],
                    [sg.Multiline('', size=(20, 5), key=self.OutputText_key)],
                    [
                        sg.Column([
                            [sg.Button('Chuyển', button_color='blue', size=(10, 1), font=("Arial Bold", 12)),
                             sg.Button('Save', button_color='blue', size=(10, 1), font=("Arial Bold", 12))]
                        ], justification='center', background_color=self.Primary_Color)
                    ],

                ], background_color=self.Primary_Color),
                sg.VerticalSeparator(pad=((20, 20), (0, 0))),

                # col 3
                sg.Column([
                    [sg.Text("Văn bản ký", background_color=self.Primary_Color)],
                    [sg.VerticalSeparator("None")],
                    [sg.Multiline('', size=(20, 5), key=self.Input)],
                    [sg.Button('File văn bản', button_color='blue', size=(10, 1), font=("Arial Bold", 12))],
                    [sg.VerticalSeparator("None")],
                    [sg.Text("Chữ ký", background_color=self.Primary_Color)],
                    [sg.Multiline('', size=(20, 5), key=self.Output)],
                    [sg.Button('File chữ ký', button_color='blue', size=(10, 1), font=("Arial Bold", 12)),
                     sg.Button('Kiểm tra', button_color='blue', font=("Arial Bold", 12))
                     ],
                ], vertical_alignment='top', background_color=self.Primary_Color),
            ],
        ]
        self.window = sg.Window('RSA Digital Signature', self.layout, size=(920, 550),
                                background_color=self.Primary_Color)
        self.is_running = True
        self.default_values = {
            self.P_key: "",
            self.Q_key: "",
            self.PublicKey: "",
            self.PrivateKey: "",
            self.Input_key: "",
            self.OutputText_key: "",
            self.Input: "",
            self.Output: "",
            self.MD5_key: "",
        }

    def run(self):
        while self.is_running:
            event, values = self.window.read()
            if event == sg.WINDOW_CLOSED:
                self.is_running = False
            elif event == 'Tạo':
                p = int(values[self.P_key])
                q = int(values[self.Q_key])
                if p <= 1 or q <= 1:
                    sg.popup('P, Q phải lơn hơn 1')
                    continue
                if not self.is_prime(p) or not self.is_prime(q):
                    sg.popup('P và Q phải là số nguyên tố')
                    continue
                e = self.public_key(p, q)
                d = self.private_key(p, q)
                n = p * q
                public_key = f'e: {e}, n: {n}'
                private_key = f'd: {d}, n: {n}'
                self.window[self.PublicKey].update(public_key)
                self.window[self.PrivateKey].update(private_key)
            elif event == 'Reset':
                for key in self.default_values:
                    self.window[key].update(self.default_values[key])
            elif event == 'File':
                file_path = sg.popup_get_file('Select a file',
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    file_content = self.read_file(file_path)
                    self.window['-Input-'].update(file_content)
            elif event == "Ký":
                self.Ky(values)
            elif event == "Chuyển":
                res = values[self.Input_key]
                tmp = values[self.OutputText_key]
                self.window['Input'].update(res)
                self.window['OutputText'].update(tmp)
            elif event == "File văn bản":
                file_path = sg.popup_get_file('Select a file',
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    file_content = self.read_file(file_path)
                    self.window['Input'].update(file_content)
            elif event == "File chữ ký":
                file_path = sg.popup_get_file('Select a file',
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    file_content = self.read_file(file_path)
                    self.window[self.Output].update(file_content)
            elif event == 'Save':

                file_path = sg.popup_get_file('Save File', save_as=True,
                                              file_types=(('Text Files', '*.txt'),
                                                          ('Word Documents', '*.docx')))
                if file_path:
                    data = values[self.OutputText_key]
                    rsa.save_data_to_file(file_path, data)
            elif event == "Kiểm tra":
                self.check_chu_ky(values)
        self.window.close()


rsa = AppRSA()
rsa.run()
